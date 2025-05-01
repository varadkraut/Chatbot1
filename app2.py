# %%
def get_docs_from_url(url):
    loader = WebBaseLoader(url)
    docs = loader.load()
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=200, chunk_overlap=20)
    split_docs = text_splitter.split_documents(docs)
    st.write('Documents Loaded from URL')
    return split_docs

# %%
from sentence_transformers import SentenceTransformer
from langchain.docstore.document import Document as LangchainDocument
import faiss
import numpy as np
import time
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument  # for .docx
from langchain.docstore.document import Document as LangchainDocument
import requests
from pathlib import Path
from langchain_text_splitters import (Language,RecursiveCharacterTextSplitter)
import os
from dotenv import load_dotenv
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.document_loaders import WebBaseLoader    
from groq import Groq


# %%
def get_docs(uploaded_file):
    start_time = time.time()
    file_name = uploaded_file.name.lower()

    # Save uploaded file temporarily
    with open("temp_file", "wb") as f:
        f.write(uploaded_file.getbuffer())

    documents = []

    if file_name.endswith(".pdf"):
        loader = PyPDFLoader("temp_file")
        documents = loader.load()

    elif file_name.endswith(".txt"):
        with open("temp_file", "r", encoding="utf-8") as f:
            content = f.read()
        documents = [LangchainDocument(page_content=content)]

    elif file_name.endswith(".docx"):
        doc = DocxDocument("temp_file")
        full_text = "\n".join([para.text for para in doc.paragraphs])
        documents = [LangchainDocument(page_content=full_text)]

    else:
        st.error("Unsupported file format. Please upload PDF, DOCX, or TXT.")
        os.remove("temp_file")
        return []

    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=100)
    final_documents = text_splitter.split_documents(documents)

    st.write("Documents Loaded")
    end_time = time.time()
    st.write(f"Time taken to load documents: {end_time - start_time:.2f} seconds")
    os.remove("temp_file")  # Clean up
    return final_documents

# %%
def create_vector_store(docs):
    start_time = time.time()
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={"trust_remote_code": True})
    vectorstore = FAISS.from_documents(docs, embeddings)
    st.write('DB is ready')
    end_time = time.time()
    st.write(f"Time taken to create DB: {end_time - start_time:.2f} seconds")
    return vectorstore

# %%
def chat_groq(messages):
    load_dotenv()
    client = Groq(api_key=os.environ.get('GROQ_API_KEY'))
    response_content = ''
    stream = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=messages,
        max_tokens=1024,
        temperature=1.3,
        stream=True,
    )

    for chunk in stream:
        content = chunk.choices[0].delta.content
        if content:
            response_content += chunk.choices[0].delta.content
    return response_content

# %%
def summarize_chat_history(chat_history):
    chat_history_text = " ".join([f"{chat['role']}: {chat['content']}" for chat in chat_history])
    prompt = f"Summarize the following chat history:\n\n{chat_history_text}"
    messages = [{'role': 'system', 'content': 'You are very good at summarizing the chat between User and Assistant'}]
    messages.append({'role': 'user', 'content': prompt})
    summary = chat_groq(messages)
    return summary

# %%
import streamlit as st

# %%
def main():
    st.set_page_config(page_title='Bajaj Finserv Chatbot')

    st.title("Bajaj Finserv Chatbot")
    with st.expander("Instructions to upload PDF, DOCX, or TXT / URL"):
        st.write("1. Pull up the side bar in top left corner.")
        st.write("2. If uploading a PDF, click 'Upload PDF', select your file, and wait for 'Documents Loaded' confirmation.")
        st.write("3. If entering a web URL, enter the URL, click 'Enter Web URL', and submit 'Process URL' and wait for 'Documents Loaded from URL' confirmation.")
        st.write("4. After loading documents, click 'Create Vector Store' to process.Documents can only be uploaded once per session")
        st.write("5. Enter a question in the text area and submit to interact with the AI chatbot.")
        st.write("6. Click on Generate Chat Summary to get the conversation of the Chat Session.")

    # Sidebar for document source selection
    st.sidebar.subheader("Choose document source:")
    option = st.sidebar.radio("Select one:", ("Upload File (PDF, DOCX, or TXT)", "Enter Web URL"))

    if "docs" not in st.session_state:
        st.session_state.docs = None
    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "user_input" not in st.session_state:
        st.session_state.user_input = ""
    if "current_prompt" not in st.session_state:
        st.session_state.current_prompt = ""
    if "chat_summary" not in st.session_state:
        st.session_state.chat_summary = ""

    if option == "Upload File (PDF, DOCX, or TXT)":
        uploaded_file = st.sidebar.file_uploader("Upload a file", type=["pdf", "docx", "txt"])

        if uploaded_file is not None:
            if st.session_state.docs is None:
                with st.spinner("Loading documents..."):
                    docs = get_docs(uploaded_file)
                st.session_state.docs = docs

    elif option == "Enter Web URL":
        url = st.sidebar.text_input("Enter URL", key="url_input")
        if st.session_state.url_input != url:
            st.session_state.url_input = url
            st.session_state.docs = None
        if st.sidebar.button('Process URL'):
            if url and st.session_state.docs is None:
                with st.spinner("Fetching and processing documents from URL..."):
                    docs = get_docs_from_url(url)
                st.session_state.docs = docs

    if st.session_state.docs is not None:
        if st.sidebar.button('Create Vector Store'):
            with st.spinner("Creating vector store..."):
                vectorstore = create_vector_store(st.session_state.docs)
            st.session_state.vectorstore = vectorstore

    if st.session_state.vectorstore is not None:
        def submit_with_doc():
            user_message = st.session_state.user_input
            if user_message:
                retriever = st.session_state.vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 3})
                context = retriever.invoke(user_message)
                prompt = f'''
                Answer the user's question based on the latest input provided in the chat history. Ignore
                previous inputs unless they are directly related to the latest question. Provide a generic
                answer if the answer to the user's question is not present in the context by mentioning it
                as general information.

                Context: {context}

                Chat History: {st.session_state.chat_history}

                Latest Question: {user_message}
                '''

                messages = [{'role': 'system', 'content': 'You are a very helpful assistant'}]
                messages.append({'role': 'user', 'content': prompt})

                try:
                    ai_response = chat_groq(messages)
                except Exception as e:
                    st.error(f"Error occurred during chat_groq execution: {str(e)}")
                    ai_response = "An error occurred while fetching response. Please try again."

                st.session_state.current_question = user_message  # Store for display

                # Display the current output prompt
                st.session_state.current_prompt = ai_response

                # Update chat history
                st.session_state.chat_history.append({'role': 'user', 'content': user_message})
                st.session_state.chat_history.append({'role': 'assistant', 'content': ai_response})

                # Clear the input field
                st.session_state.user_input = ""

    def submit_without_doc():
        user_message = st.session_state.user_input
        if user_message:
            prompt = f'''
            Answer the user's question based on the latest input provided in the chat history. Ignore
            previous inputs unless they are directly related to the latest
            question. 
            
            Chat History: {st.session_state.chat_history}

            Latest Question: {user_message}
            '''

            messages = [{'role': 'system', 'content': 'You are a very helpful assistant'}]
            messages.append({'role': 'user', 'content': prompt})

            try:
                ai_response = chat_groq(messages)
            except Exception as e:
                st.error(f"Error occurred during chat_groq execution: {str(e)}")
                ai_response = "An error occurred while fetching response. Please try again."

            st.session_state.current_question = user_message  # Store for display

            # Display the current output prompt
            st.session_state.current_prompt = ai_response

            # Update chat history
            st.session_state.chat_history.append({'role': 'user', 'content': user_message})
            st.session_state.chat_history.append({'role': 'assistant', 'content': ai_response})

            # Clear the input field
            st.session_state.user_input = ""

    

    # Display full chat history using Streamlit's chat_message
    if "chat_history" in st.session_state:
        for chat in st.session_state.chat_history:
            with st.chat_message(chat["role"]):
                st.markdown(chat["content"])

    # Real-time input using st.chat_input
    user_message = st.chat_input("Ask your question here...")

    if user_message:
        with st.spinner("Assistant is typing..."):
            if st.session_state.vectorstore is not None:
                # Use documents
                retriever = st.session_state.vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 3})
                context = retriever.invoke(user_message)
                prompt = f'''
                Answer the user's question based on the latest input provided in the chat history. Ignore
                previous inputs unless they are directly related to the latest question. Provide a generic
                answer if the answer to the user's question is not present in the context by mentioning it
                as general information.

                Context: {context}

                Chat History: {st.session_state.chat_history}

                Latest Question: {user_message}
                '''
            else:
                # No documents
                prompt = f'''
                Answer the user's question based on the latest input provided in the chat history. Ignore
                previous inputs unless they are directly related to the latest question.

                Chat History: {st.session_state.chat_history}

                Latest Question: {user_message}
                '''

            messages = [{'role': 'system', 'content': 'You are a very helpful assistant'}]
            messages.append({'role': 'user', 'content': prompt})

            try:
                ai_response = chat_groq(messages)
            except Exception as e:
                st.error(f"Error occurred during chat_groq execution: {str(e)}")
                ai_response = "An error occurred while fetching response. Please try again."

            # Update chat history
            st.session_state.chat_history.append({'role': 'user', 'content': user_message})
            st.session_state.chat_history.append({'role': 'assistant', 'content': ai_response})

            # Show current user message and response in chat bubbles
            with st.chat_message("user"):
                st.markdown(user_message)
            with st.chat_message("assistant"):
                st.markdown(ai_response)
    # if "current_question" in st.session_state and st.session_state.current_question:
    #     st.markdown(f"**Current Question:** {st.session_state.current_question}")


    # # Display user's current question in chat format
    # if "current_question" in st.session_state and st.session_state.current_question:
    #     with st.chat_message("user"):
    #         st.markdown(st.session_state.current_question)

    # # Display assistant's response in chat format
    # if "current_prompt" in st.session_state and st.session_state.current_prompt:
    #     with st.chat_message("assistant"):
    #         st.markdown(st.session_state.current_prompt)


    # # Button to generate chat summary
    # if st.button('Generate Chat Summary'):
    #     st.session_state.chat_summary = summarize_chat_history(st.session_state.chat_history)

    # # Display the chat summary if available
    # if st.session_state.chat_summary:
    #     with st.expander("Chat Summary"):
    #         st.write(st.session_state.chat_summary)

    # # Display the last 4 messages in an expander
    # with st.expander("Recent Chat History"):
    #     recent_history = st.session_state.chat_history[-8:][::-1]
    #     reversed_history = []
    #     for i in range(0, len(recent_history), 2):
    #         if i+1 < len(recent_history):
    #             reversed_history.extend([recent_history[i+1], recent_history[i]])
    #         else:
    #             reversed_history.append(recent_history[i])
    #     for chat in reversed_history:
    #         st.write(f"{chat['role'].capitalize()}: {chat['content']}")

if __name__ == "__main__":
    main()


